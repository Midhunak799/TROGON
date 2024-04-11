import json
from docx import Document



def word_document(file_path,question):
    try:
        doc = Document()
        
        for i,q in enumerate(question ,start=1):
            doc.add_paragraph(f"{i}. {q['question']}")
            doc.add_paragraph(f"A. {q['options']['A']}")
            doc.add_paragraph(f"B. {q['options']['B']}")
            doc.add_paragraph(f"C. {q['options']['C']}")
            doc.add_paragraph(f"D. {q['options']['D']}")
            doc.add_paragraph(f"Anwer. {q['answer']}")
            
        doc.save(file_path)
        print("doc created successfully",file_path)
        return file_path
    
    except Exception as e:
        print("error",e)
        return None


questions = [
    {
        "question": "What is the capital of France?",
        "options": {
            "A": "Berlin",
            "B": "Madrid",
            "C": "Paris",
            "D": "Lisbon"
        },
        "answer": "C"
    },
    {
        "question": "Who wrote Hamlet?",
        "options": {
            "A": "Mark Twain",
            "B": "William Shakespeare",
            "C": "Jane Austen",
            "D": "Charles Dickens"
        },
        "answer": "B"
    },
    
    {
        "question": "Who wrote Ramayana?",
        "options": {
            "A": "Valmiki",
            "B": "Jk Rowling",
            "C": "Jane Austen",
            "D": "Paulo Coelo"
        },
        "answer": "A"
    },
    
]

word_document("mcqs.docx",questions)




def to_json(file_path):
    doc = Document(file_path)
    data = [p.text for p in doc.paragraphs]
    json_data = {"data": data}
    return json_data

file_path = "mcqs.docx"
json_data = to_json(file_path)


with open("output.json", "w") as json_file:
    json.dump(json_data, json_file, indent=4)




    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
